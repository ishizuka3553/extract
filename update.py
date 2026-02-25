import docx
import time
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

def _no_wrap(text: str) -> str:
    if text is None:
        return ""
    return text.replace(" ", "\u00A0")

def _fit_text(text: str, tbl, row_idx: int, col_idx: int, font_size: int = 11) -> str:
    if text is None:
        return ""
    # remove newlines and convert spaces to non-breaking
    t = _no_wrap(text.replace("\n", " "))
    try:
        col = tbl.columns[col_idx]
        width_emu = getattr(col, 'width', None)
        if width_emu and width_emu > 0:
            width_pt = width_emu / 12700.0
            avg_char_pt = max(1.0, font_size * 0.5)
            max_chars = max(1, int(width_pt / avg_char_pt))
        else:
            max_chars = 30
    except Exception:
        max_chars = 30
    if len(t) <= max_chars:
        return t
    return t[:max_chars]

def parse_docx_tables(docx_path):
    doc = docx.Document(docx_path)
    count = 1
    for tbl in doc.tables:
        if len(tbl.rows) < 2:
            continue
        print(f"{count}ページ目を読み取り中")
        count += 1
        for row in range(2, len(tbl.rows)):
            if row < 2:
                continue
            if tbl.cell(row, 0).text.strip() == "":
                continue
            if row % 2 == 1:
                last_date(tbl, row)
    doc.save("更新後.docx")

def last_date(tbl, row):
    last_col = 2
    name = ""
    for col in range(2, len(tbl.columns)):
        if tbl.cell(row, col).text.strip() == "":
            last_col = col - 1
            name = tbl.cell(row - 1, last_col).text
            break
    if last_col % 2 == 0:
        tbl.cell(row - 1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tbl.cell(row - 1, 2).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        tbl.cell(row - 1, 2).text = name
        tbl.cell(row, 2).text = tbl.cell(row, last_col).text
        tbl.cell(row - 1, 1).text = tbl.cell(row, last_col - 1).text
        for col in range(3, last_col + 1):
            if col % 2 == 0:
                tbl.cell(row - 1, col).text = ""
            tbl.cell(row, col).text = ""
    else:
        tbl.cell(row, 1).text = tbl.cell(row, last_col).text
        for col in range(2, last_col + 1):
            if col % 2 == 0:
                tbl.cell(row - 1, col).text = ""
            tbl.cell(row, col).text = ""      

def main():
    start = time.time()
    parse_docx_tables("北茨城・高萩区域2025.docx")
    end = time.time()
    print(f"実行時間(秒):{end - start}")

if __name__ == "__main__":
    main()